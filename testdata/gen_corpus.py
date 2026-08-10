#!/usr/bin/env python3
"""
lawdesk 테스트 코퍼스 생성기.

실제 법률사무소 문서 폴더와 비슷한 구조/포맷/내용을 합성한다.
 - DOCX  : python-docx (진짜 OOXML — 제목 스타일, 표 포함)
 - PDF   : reportlab + 나눔 TTF 임베드 (한글 CID 폰트 = 추출 난이도 높은 케이스)
 - 스캔PDF: PIL로 이미지 렌더 → 기울임/노이즈/JPEG 열화/도장 → PDF 삽입

여기 들어있는 인명·회사명·사건번호는 전부 가공의 것이다.
"""
import io
import os
import random
import shutil
import sys

from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

FONT_DIR = "/usr/share/fonts/truetype/nanum"
GOTHIC = os.path.join(FONT_DIR, "NanumGothic.ttf")
MYEONGJO = os.path.join(FONT_DIR, "NanumMyeongjo.ttf")
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "corpus")

pdfmetrics.registerFont(TTFont("NanumGothic", GOTHIC))
pdfmetrics.registerFont(TTFont("NanumMyeongjo", MYEONGJO))

random.seed(20260810)  # 재현 가능하게


# ----------------------------------------------------------------- DOCX

def docx_yongyeok(path):
    d = Document()
    d.add_heading("소프트웨어 개발 용역계약서", level=0)
    d.add_paragraph(
        "주식회사 가나테크(이하 “갑”이라 한다)와 주식회사 다라소프트(이하 “을”이라 한다)는 "
        "소프트웨어 개발 용역과 관련하여 다음과 같이 계약을 체결한다."
    )

    d.add_heading("제1조 (계약의 목적)", level=1)
    d.add_paragraph(
        "본 계약은 갑이 을에게 위탁하는 통합 물류관리 시스템 개발 용역의 수행에 관하여 "
        "양 당사자의 권리와 의무를 정함을 목적으로 한다."
    )

    d.add_heading("제2조 (계약기간)", level=1)
    d.add_paragraph("계약기간은 2023년 4월 1일부터 2024년 3월 31일까지로 한다.")

    d.add_heading("제3조 (계약금액 및 지급방법)", level=1)
    d.add_paragraph(
        "① 총 계약금액은 금 일억오천만원(￦150,000,000, 부가가치세 별도)으로 한다."
    )
    d.add_paragraph("② 갑은 다음 각 호에 따라 대금을 지급한다.")

    t = d.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    rows = [
        ("구분", "지급시기", "금액"),
        ("착수금", "계약 체결일로부터 7일 이내", "45,000,000원"),
        ("중도금", "중간 산출물 검수 완료 시", "60,000,000원"),
        ("잔금", "최종 검수 완료 후 30일 이내", "45,000,000원"),
    ]
    for i, r in enumerate(rows):
        for j, v in enumerate(r):
            t.cell(i, j).text = v

    d.add_heading("제7조 (손해배상)", level=1)
    d.add_paragraph(
        "① 을의 귀책사유로 인하여 갑에게 손해가 발생한 경우 을은 그 손해를 배상할 책임을 진다."
    )
    d.add_paragraph(
        "② 전항의 배상책임은 특별한 사정이 없는 한 총 계약금액을 한도로 한다. "
        "다만 을의 고의 또는 중대한 과실로 인한 경우에는 그러하지 아니하다."
    )

    d.add_heading("제12조 (관할법원)", level=1)
    d.add_paragraph(
        "본 계약과 관련하여 분쟁이 발생한 경우 서울중앙지방법원을 제1심 관할법원으로 한다."
    )
    d.save(path)


def docx_imdaecha(path):
    d = Document()
    d.add_heading("부동산 임대차계약서", level=0)

    d.add_heading("제1조 (목적물의 표시)", level=1)
    d.add_paragraph("소재지: 서울특별시 강남구 테헤란로 123, 5층 501호")
    d.add_paragraph("면적: 전용면적 132.23㎡ (약 40평)")
    d.add_paragraph("용도: 사무실")

    d.add_heading("제2조 (보증금 및 차임)", level=1)
    d.add_paragraph("① 보증금은 금 오억원(￦500,000,000)으로 한다.")
    d.add_paragraph("② 월 차임은 금 삼백오십만원(￦3,500,000)으로 하며, 매월 말일에 선불로 지급한다.")
    d.add_paragraph("③ 관리비는 차임에 포함되지 아니하며 실비로 정산한다.")

    d.add_heading("제5조 (원상회복의무)", level=1)
    d.add_paragraph(
        "임차인은 계약이 종료된 때에는 목적물을 원상으로 회복하여 임대인에게 반환하여야 한다. "
        "다만 임대인의 동의를 얻어 부속시킨 물건에 대하여는 매수를 청구할 수 있다."
    )

    d.add_heading("제9조 (계약의 해지)", level=1)
    d.add_paragraph(
        "임차인이 차임을 3기 이상 연체한 경우 임대인은 최고 없이 본 계약을 해지할 수 있다."
    )
    d.save(path)


def docx_nda(path):
    d = Document()
    d.add_heading("비밀유지계약서 (Non-Disclosure Agreement)", level=0)
    d.add_paragraph(
        "본 계약은 대한민국 법률에 따라 규율되며 해석된다. "
        "This Agreement shall be governed by and construed in accordance with "
        "the laws of the Republic of Korea."
    )

    d.add_heading("제2조 (비밀정보의 정의)", level=1)
    d.add_paragraph(
        "“비밀정보”라 함은 서면, 구두, 전자적 형태 기타 방법으로 제공되는 일체의 기술상·경영상 정보로서 "
        "비밀로 표시되었거나 그 성질상 비밀로 취급함이 상당한 정보를 말한다."
    )

    d.add_heading("제4조 (비밀유지의무)", level=1)
    d.add_paragraph(
        "수령당사자는 비밀정보를 선량한 관리자의 주의의무로 보호하여야 하며, "
        "제공당사자의 사전 서면 동의 없이 제3자에게 누설하거나 본 계약의 목적 외로 사용하여서는 아니 된다."
    )
    d.add_paragraph(
        "Liability for damages arising out of a breach of this Article shall not exceed "
        "the total contract amount, except in cases of willful misconduct or gross negligence."
    )

    d.add_heading("제6조 (계약기간 및 존속)", level=1)
    d.add_paragraph(
        "본 계약의 유효기간은 체결일로부터 3년으로 하며, 비밀유지의무는 계약 종료 후 5년간 존속한다."
    )
    d.save(path)


# ------------------------------------------------------------ 텍스트 PDF

def draw_pdf(path, title, lines, font="NanumMyeongjo"):
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    y = h - 30 * mm

    c.setFont("NanumGothic", 15)
    c.drawCentredString(w / 2, y, title)
    y -= 14 * mm

    c.setFont(font, 10.5)
    for line in lines:
        if y < 25 * mm:
            c.showPage()
            c.setFont(font, 10.5)
            y = h - 25 * mm
        if line == "":
            y -= 5 * mm
            continue
        c.drawString(25 * mm, y, line)
        y -= 6.2 * mm
    c.save()


PANRYE_1 = [
    "【사건】 2021다12345 손해배상(기)",
    "【원고, 상고인】 김○○",
    "【피고, 피상고인】 주식회사 마바건설",
    "【원심판결】 서울고등법원 2021. 1. 20. 선고 2020나56789 판결",
    "【선고일】 2021. 5. 13.",
    "",
    "【주 문】",
    "원심판결을 파기하고, 사건을 서울고등법원에 환송한다.",
    "",
    "【이 유】",
    "상고이유를 판단한다.",
    "",
    "1. 채무불이행으로 인한 손해배상책임의 성립에 관하여",
    "",
    "  가. 채무자가 채무의 내용에 좇은 이행을 하지 아니한 때에는 채권자는 손해배상을",
    "청구할 수 있다(민법 제390조). 다만 채무자의 고의나 과실 없이 이행할 수 없게 된",
    "때에는 그러하지 아니하다.",
    "",
    "  나. 원심은 피고가 공사기간을 준수하지 못한 사실은 인정하면서도, 자재 수급 지연이",
    "불가항력에 해당한다는 이유로 피고의 손해배상책임을 부정하였다.",
    "",
    "  다. 그러나 기록에 의하면 자재 수급 지연은 피고가 사전에 예견할 수 있었던 사정으로",
    "보이고, 피고가 이를 회피하기 위한 상당한 조치를 취하였다고 볼 자료가 없다.",
    "따라서 이를 불가항력으로 단정한 원심의 판단에는 채무불이행에 있어 귀책사유에 관한",
    "법리를 오해하여 판결에 영향을 미친 잘못이 있다.",
    "",
    "2. 손해액의 산정에 관하여",
    "",
    "  손해배상액은 통상의 손해를 한도로 하되, 특별한 사정으로 인한 손해는 채무자가 그",
    "사정을 알았거나 알 수 있었을 때에 한하여 배상책임이 있다(민법 제393조).",
    "원고가 주장하는 영업손실 5,000만원 부분은 특별손해에 해당하므로, 환송 후 원심은",
    "피고의 예견가능성에 관하여 심리하여야 한다.",
    "",
    "3. 결론",
    "",
    "그러므로 원심판결을 파기하고 사건을 원심법원에 환송하기로 하여 관여 대법관의",
    "일치된 의견으로 주문과 같이 판결한다.",
]

PANRYE_2 = [
    "【사건】 2023고합987 특정경제범죄가중처벌등에관한법률위반(횡령)",
    "【피고인】 이○○ (19○○년생, 회사원)",
    "【검사】 박○○",
    "【변호인】 변호사 최○○",
    "【선고일】 2023. 11. 2.",
    "",
    "【주 문】",
    "피고인을 징역 1년에 처한다.",
    "다만, 이 판결 확정일로부터 2년간 위 형의 집행을 유예한다.",
    "피고인에게 120시간의 사회봉사를 명한다.",
    "",
    "【범죄사실】",
    "피고인은 2022. 3.경부터 2022. 12.경까지 주식회사 사아상사의 자금관리 업무에",
    "종사하면서, 업무상 보관하던 회사 자금 합계 3억 2,000만원을 개인 채무 변제 등에",
    "임의로 사용하여 이를 횡령하였다.",
    "",
    "【증거의 요지】",
    "1. 피고인의 법정진술",
    "1. 증인 정○○의 법정진술",
    "1. 계좌거래내역서, 회계장부 사본",
    "",
    "【양형의 이유】",
    "피고인이 범행을 인정하고 깊이 반성하고 있는 점, 피해액 전액을 변제하여 피해",
    "회사와 원만히 합의한 점, 동종 전과가 없는 초범인 점 등을 유리한 정상으로 참작한다.",
    "다만 횡령액이 적지 아니하고 범행이 상당 기간 반복된 점은 불리한 정상이다.",
    "이러한 사정들을 종합하여 주문과 같이 형을 정한다.",
]


# ---------------------------------------------------------- 스캔 이미지 PDF

def render_scan_image(title, lines, seal=True, quality=45, angle=0.45):
    """종이 문서를 스캐너로 찍은 듯한 이미지를 만든다."""
    W, H = 1240, 1754  # A4 @150dpi
    img = Image.new("L", (W, H), 252)
    dr = ImageDraw.Draw(img)

    f_title = ImageFont.truetype(GOTHIC, 40)
    f_body = ImageFont.truetype(MYEONGJO, 27)

    dr.text((W // 2, 120), title, font=f_title, fill=20, anchor="mm")
    y = 220
    for line in lines:
        if line == "":
            y += 20
            continue
        dr.text((110, y), line, font=f_body, fill=random.randint(15, 55))
        y += 46
        if y > H - 150:
            break

    if seal:
        # 붉은 직인 — OCR 방해 요소
        sx, sy, r = W - 260, H - 320, 85
        dr.ellipse([sx - r, sy - r, sx + r, sy + r], outline=90, width=6)
        f_seal = ImageFont.truetype(GOTHIC, 30)
        dr.text((sx, sy), "인", font=f_seal, fill=90, anchor="mm")

    # 스캔 열화 시뮬레이션
    img = img.rotate(angle, resample=Image.BICUBIC, fillcolor=250, expand=False)
    img = img.filter(ImageFilter.GaussianBlur(0.6))
    px = img.load()
    for _ in range(int(W * H * 0.012)):  # 점 노이즈
        x, yy = random.randrange(W), random.randrange(H)
        px[x, yy] = max(0, min(255, px[x, yy] + random.randint(-70, 45)))

    buf = io.BytesIO()
    img.convert("L").save(buf, format="JPEG", quality=quality)
    buf.seek(0)
    return buf


def scan_pdf(path, title, lines, **kw):
    buf = render_scan_image(title, lines, **kw)
    from reportlab.lib.utils import ImageReader
    c = canvas.Canvas(path, pagesize=A4)
    c.drawImage(ImageReader(buf), 0, 0, width=A4[0], height=A4[1])
    c.save()


DEUNGGI = [
    "[집합건물] 서울특별시 강남구 테헤란로 123",
    "",
    "【표 제 부】 (전유부분의 건물의 표시)",
    "표시번호   접수           건물번호      건물내역",
    "1         2015년 3월 2일   제5층 제501호  철근콘크리트조 132.23㎡",
    "",
    "【갑    구】 (소유권에 관한 사항)",
    "순위번호  등기목적    접수              등기원인      권리자",
    "1        소유권보존   2015년 3월 2일     -            주식회사 가나테크",
    "2        소유권이전   2019년 7월 15일    2019년 6월 3일 매매",
    "                                                   주식회사 마바건설",
    "",
    "【을    구】 (소유권 이외의 권리에 관한 사항)",
    "순위번호  등기목적    접수              등기원인      권리자",
    "1        근저당권설정  2019년 7월 15일   설정계약      채권최고액 금 육억원",
    "                                                   근저당권자 주식회사 국민은행",
]

JEUNGGEO = [
    "물품공급계약서 (사본)",
    "",
    "공급자: 주식회사 사아상사",
    "수요자: 주식회사 자차유통",
    "",
    "제1조 공급자는 수요자에게 별지 목록 기재 물품을 공급한다.",
    "제2조 대금은 총 금 삼억이천만원으로 하며, 월 단위로 정산한다.",
    "제3조 수요자는 물품 인수일로부터 30일 이내에 대금을 지급한다.",
    "제4조 대금 지급이 지연되는 경우 연 12%의 지연손해금을 가산한다.",
    "",
    "2022년 3월 10일",
    "",
    "공급자  주식회사 사아상사  대표이사 정○○",
    "수요자  주식회사 자차유통  대표이사 이○○",
]


# ---------------------------------------------------------------- 메인

def main():
    if os.path.isdir(OUT):
        shutil.rmtree(OUT)
    for sub in ("계약서", "판례", "증거자료", "메모"):
        os.makedirs(os.path.join(OUT, sub), exist_ok=True)

    made = []

    def mark(p):
        made.append(p)
        return p

    docx_yongyeok(mark(f"{OUT}/계약서/용역계약서_주식회사가나테크_2023.docx"))
    docx_imdaecha(mark(f"{OUT}/계약서/임대차계약서_테헤란로501호.docx"))
    docx_nda(mark(f"{OUT}/계약서/비밀유지계약서_NDA.docx"))

    draw_pdf(mark(f"{OUT}/판례/대법원_2021다12345_손해배상.pdf"),
             "대 법 원 판 결", PANRYE_1)
    draw_pdf(mark(f"{OUT}/판례/서울중앙지방법원_2023고합987_횡령.pdf"),
             "서울중앙지방법원 판결", PANRYE_2)

    scan_pdf(mark(f"{OUT}/증거자료/등기부등본_강남테헤란로_스캔.pdf"),
             "등기사항전부증명서", DEUNGGI, seal=True, quality=45, angle=0.45)
    scan_pdf(mark(f"{OUT}/증거자료/갑제3호증_물품공급계약서사본.pdf"),
             "물품공급계약서", JEUNGGEO, seal=True, quality=30, angle=-0.9)

    with open(mark(f"{OUT}/메모/사건메모_2021다12345.txt"), "w", encoding="utf-8") as f:
        f.write(
            "2021다12345 손해배상 사건 메모\n\n"
            "- 쟁점: 자재 수급 지연이 불가항력에 해당하는지\n"
            "- 원심은 불가항력 인정 → 대법원 파기환송\n"
            "- 영업손실 5,000만원은 특별손해, 예견가능성 심리 필요\n"
            "- 관련 조문: 민법 제390조, 제393조\n"
        )
    with open(mark(f"{OUT}/메모/체크리스트.md"), "w", encoding="utf-8") as f:
        f.write(
            "# 계약서 검토 체크리스트\n\n"
            "- [ ] 손해배상 한도 조항 유무\n"
            "- [ ] 관할법원 합의 조항\n"
            "- [ ] 비밀유지의무 존속기간\n"
            "- [ ] 원상회복의무 범위\n"
        )

    print(f"생성 완료: {len(made)}개 파일\n")
    for p in sorted(made):
        rel = os.path.relpath(p, OUT)
        print(f"  {os.path.getsize(p):>9,}B  {rel}")


if __name__ == "__main__":
    sys.exit(main())
